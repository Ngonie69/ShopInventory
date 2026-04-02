"""Generate End-User Guide Word Document for Kefalos Portal"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from datetime import datetime

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)

def add_tip(text):
    p = doc.add_paragraph()
    run = p.add_run('💡 Tip: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x0d, 0x6e, 0xfd)
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run('📌 Note: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xdc, 0x35, 0x45)
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(f' – {text}')
    else:
        p.add_run(text)

def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for par in cell.paragraphs:
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('KEFALOS PORTAL')
r.font.size = Pt(36)
r.bold = True
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('End-User Guide')
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Kefalos Quality Dairy Produce')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Inventory Management & SAP Integration System')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Version 1.0  |  {datetime.now().strftime("%B %Y")}')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CONFIDENTIAL – For Internal Use Only')
r.font.size = Pt(10)
r.bold = True
r.font.color.rgb = RGBColor(0xdc, 0x35, 0x45)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction',
    '2. Getting Started',
    '   2.1 Logging In',
    '   2.2 Navigation Overview',
    '   2.3 Dashboard',
    '3. Sales & Billing',
    '   3.1 Invoices',
    '   3.2 Sales Orders',
    '   3.3 Quotations',
    '   3.4 Credit Notes',
    '4. Inventory Management',
    '   4.1 Products',
    '   4.2 Stock Monitoring',
    '   4.3 Prices',
    '   4.4 Inventory Transfers',
    '5. Purchasing',
    '   5.1 Purchase Orders',
    '6. Payments',
    '   6.1 Viewing Payments',
    '   6.2 Recording a Payment',
    '7. Customers',
    '   7.1 Customer List',
    '   7.2 G/L Accounts',
    '8. Insights & Reporting',
    '   8.1 Reports',
    '   8.2 Proof of Delivery (POD)',
    '   8.3 POD Dashboard',
    '   8.4 POD Upload Report',
    '9. Documents & Templates',
    '   9.1 Document Templates',
    '   9.2 Document Manager',
    '10. System Tools',
    '   10.1 Notifications',
    '   10.2 Exchange Rates',
    '   10.3 Sync Status',
    '   10.4 REVMax Fiscal Device',
    '   10.5 AI Assistant',
    '11. Administration (Admin Only)',
    '   11.1 User Management',
    '   11.2 Security Settings',
    '   11.3 Audit Trail',
    '   11.4 User Activity',
    '   11.5 Backups',
    '   11.6 Settings',
    '   11.7 Customer Portal Management',
    '   11.8 Webhooks',
    '   11.9 API Explorer',
    '12. Customer Portal (External)',
    '   12.1 Customer Login',
    '   12.2 Customer Dashboard',
    '   12.3 Customer Invoices',
    '   12.4 Customer Statements',
    '   12.5 Customer Payments',
    '   12.6 Customer PODs',
    '   12.7 Customer Support',
    '13. Troubleshooting & FAQ',
    'Appendix A: User Roles & Permissions',
    'Appendix B: Keyboard Shortcuts',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('1. Introduction', level=1)

doc.add_paragraph(
    'Welcome to the Kefalos Portal, the centralised web application for managing '
    'inventory, sales, purchasing, and customer operations at Kefalos Quality Dairy Produce. '
    'This system integrates with SAP Business One to provide real-time data across all '
    'business functions.'
)

doc.add_paragraph(
    'This guide covers every feature available to staff users and provides step-by-step '
    'instructions for common workflows. A separate section covers the Customer Portal '
    'for external customers.'
)

doc.add_heading('What You Can Do', level=3)
add_bullet('Create and manage invoices, sales orders, quotations, and credit notes')
add_bullet('Monitor stock levels across all warehouses in real time')
add_bullet('Record customer payments and track outstanding balances')
add_bullet('Generate business reports and export to Excel or PDF')
add_bullet('Upload and track proof of delivery (POD) for shipments')
add_bullet('Manage customers, products, and pricing')
add_bullet('Fiscalise transactions via REVMax fiscal device')
add_bullet('Access the AI Assistant for quick data insights')

doc.add_heading('System Requirements', level=3)
add_bullet('Modern web browser (Chrome, Edge, Firefox, or Safari)')
add_bullet('Internet/network access to the company server')
add_bullet('Valid user account with assigned role')
add_bullet('The application supports dark mode and works on desktop and mobile devices')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 2. GETTING STARTED
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('2. Getting Started', level=1)

# 2.1 Logging In
doc.add_heading('2.1 Logging In', level=2)
doc.add_paragraph('To access the Kefalos Portal:')
add_numbered('Open your web browser and navigate to the portal URL provided by your administrator.')
add_numbered('Enter your Username and Password.')
add_numbered('Click the "Sign In" button.')
add_numbered('If two-factor authentication (2FA) is enabled, enter the verification code sent to your email or authenticator app.')

add_tip('Use the eye icon next to the password field to show or hide your password while typing.')
add_note('If you have forgotten your password, contact your system administrator to reset it.')

doc.add_paragraph(
    'After a successful login you will be taken to the Dashboard. Your session will remain '
    'active until you log out or the session times out due to inactivity.'
)

# 2.2 Navigation Overview
doc.add_heading('2.2 Navigation Overview', level=2)
doc.add_paragraph(
    'The application uses a sidebar navigation menu on the left and a top navigation bar. '
    'The sidebar is organised into the following sections:'
)

add_table(
    ['Section', 'Pages', 'Who Can See It'],
    [
        ['Overview', 'Dashboard', 'All staff roles'],
        ['Sales & Billing', 'Invoices, Sales Orders, Quotations, Credit Notes', 'Admin, Cashier'],
        ['Inventory', 'Stock, Products, Inventory Transfers', 'Admin, Cashier, StockController, DepotController, Manager'],
        ['Purchasing', 'Purchase Orders', 'Admin, Manager'],
        ['Payments', 'Payments', 'Admin, Cashier, DepotController'],
        ['Catalogue', 'Customers, Prices, G/L Accounts, Exchange Rates', 'All staff roles (Prices: Admin only)'],
        ['Insights', 'Reports, POD Dashboard, Proof of Delivery, POD Report', 'Varies by role'],
        ['System', 'Notifications, Sync Status, REVMax, AI Assistant', 'Varies by role'],
        ['Administration', 'User Management, Settings, Audit Trail, Backups, etc.', 'Admin only'],
    ]
)

doc.add_paragraph('The top navigation bar provides:')
add_bullet('Notification bell', 'Quick access to alerts')
add_bullet('Theme toggle', 'Switch between light and dark mode')
add_bullet('Profile menu', 'Access settings, security, and logout')

add_tip('On mobile devices, tap the hamburger menu icon (☰) to open the sidebar.')

# 2.3 Dashboard
doc.add_heading('2.3 Dashboard', level=2)
doc.add_paragraph(
    'The Dashboard is your home page and shows a real-time overview of key business metrics:'
)
add_bullet('Total Invoices – click to go to the Invoices page')
add_bullet('Total Payments – click to go to the Payments page')
add_bullet('Live status indicator showing system connectivity')
add_bullet('Quick-action links to frequently used modules')

add_tip('The dashboard refreshes every time you navigate to it. Pull down on mobile to refresh.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 3. SALES & BILLING
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('3. Sales & Billing', level=1)
doc.add_paragraph(
    'The Sales & Billing section lets you create and manage all customer-facing documents: '
    'invoices, sales orders, quotations, and credit notes. All documents are synchronised '
    'with SAP Business One.'
)

# 3.1 Invoices
doc.add_heading('3.1 Invoices', level=2)
doc.add_paragraph(
    'The Invoices page displays all customer invoices with summary statistics and powerful '
    'filtering options.'
)

doc.add_heading('Viewing Invoices', level=3)
add_bullet('Summary cards at the top show: Total Invoices, Paid, Unpaid, and Overdue counts')
add_bullet('Use the search bar to find invoices by number or customer name')
add_bullet('Filter by date range, customer, status, or payment terms')
add_bullet('Click any invoice row to open the detail drawer on the right')

doc.add_heading('Invoice Statuses', level=3)
add_table(
    ['Status', 'Meaning'],
    [
        ['Draft', 'Invoice saved but not yet posted to SAP'],
        ['Pending', 'Invoice posted, awaiting payment'],
        ['Sent', 'Invoice emailed to the customer'],
        ['Paid', 'Full payment received'],
        ['Overdue', 'Payment due date has passed'],
    ]
)

doc.add_heading('Creating a New Invoice', level=3)
add_numbered('Click the "Create Invoice" button or navigate to Sales & Billing → Create Invoice.')
add_numbered('Select (or search for) the customer from the dropdown.')
add_numbered('Choose a Price List and Currency.')
add_numbered('Add line items by searching for products. Enter the quantity for each item.')
add_numbered('Prices are calculated automatically based on the selected price list. You can override the price if needed.')
add_numbered('Review the totals (Net Total, VAT, Invoice Total) at the bottom.')
add_numbered('Click "Save" to post the invoice to SAP, or "Save & Print" to save and generate a PDF.')

doc.add_heading('Exporting Invoices', level=3)
add_bullet('Export to Excel – downloads the filtered invoice list as a spreadsheet')
add_bullet('Export to PDF – generates a PDF of the currently displayed invoices')
add_bullet('Email – send a specific invoice directly to the customer\'s email')

add_tip('Use the date range filter to narrow down invoices before exporting for cleaner reports.')

# 3.2 Sales Orders
doc.add_heading('3.2 Sales Orders', level=2)
doc.add_paragraph(
    'Sales Orders represent confirmed customer orders that have not yet been invoiced. '
    'They can be created independently or converted from approved quotations.'
)

doc.add_heading('Viewing Sales Orders', level=3)
add_bullet('Summary cards show: Total, Pending, Processing, Delivered, and Cancelled orders')
add_bullet('Filter by customer, date range, or status')
add_bullet('Click any order to view full details including line items')

doc.add_heading('Creating a Sales Order', level=3)
add_numbered('Navigate to Sales & Billing → Create Sale Order.')
add_numbered('Select the customer.')
add_numbered('Add products with quantities.')
add_numbered('Add any special instructions in the Notes field.')
add_numbered('Click "Save" to create the order in SAP.')

doc.add_heading('Converting to Invoice', level=3)
doc.add_paragraph(
    'To convert a sales order into an invoice, open the order details and click "Convert to Invoice". '
    'This copies all line items into a new invoice document.'
)
add_note('Once converted, the sales order status will change to reflect the invoice creation.')

# 3.3 Quotations
doc.add_heading('3.3 Quotations', level=2)
doc.add_paragraph(
    'Quotations are price estimates sent to customers before they confirm an order. They have '
    'an expiry date and can be converted to sales orders or invoices once accepted.'
)

doc.add_heading('Quotation Statuses', level=3)
add_table(
    ['Status', 'Meaning'],
    [
        ['Draft', 'Quotation being prepared'],
        ['Sent', 'Quotation sent to customer'],
        ['Accepted', 'Customer accepted the quote'],
        ['Rejected', 'Customer declined the quote'],
        ['Expired', 'Quotation validity period has passed'],
        ['Converted', 'Converted to a sales order or invoice'],
    ]
)

doc.add_heading('Creating a Quotation', level=3)
add_numbered('Navigate to Sales & Billing → Create Quotation.')
add_numbered('Select the customer and price list.')
add_numbered('Add line items with products and quantities.')
add_numbered('Set the quotation expiry date.')
add_numbered('Add any terms or notes.')
add_numbered('Click "Save" then use "Email" to send to the customer.')

add_tip('Once a customer accepts a quotation, open it and click "Convert to Sales Order" or "Convert to Invoice".')

# 3.4 Credit Notes
doc.add_heading('3.4 Credit Notes', level=2)
doc.add_paragraph(
    'Credit Notes are issued to adjust or reverse all or part of an invoice, for example '
    'due to returned goods, pricing errors, or agreed discounts.'
)

doc.add_heading('Creating a Credit Note', level=3)
add_numbered('Navigate to Sales & Billing → Create Credit Note.')
add_numbered('Reference the original invoice number.')
add_numbered('Select a reason: Damage, Overcharge, Return, Discount, etc.')
add_numbered('Add the line items to credit (partial or full reversal).')
add_numbered('Add any explanatory notes.')
add_numbered('Click "Save" to post the credit note to SAP.')

add_note('Credit notes affect the customer\'s account balance and are visible in the Customer Portal.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 4. INVENTORY MANAGEMENT
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('4. Inventory Management', level=1)

# 4.1 Products
doc.add_heading('4.1 Products', level=2)
doc.add_paragraph(
    'The Products page is your product master database. It shows all items synced from SAP '
    'Business One with their current status and details.'
)

add_bullet('Summary cards show: Total Products, Active SKUs, Batches, and Warehouses')
add_bullet('Search by product code, name, or barcode')
add_bullet('Filter by category, status (Active, Discontinued), or warehouse')
add_bullet('Click any product to view details: code, name, unit of measure, category, and pricing')

add_tip('Use the search bar with a barcode scanner for quick product lookup.')

# 4.2 Stock Monitoring
doc.add_heading('4.2 Stock Monitoring', level=2)
doc.add_paragraph(
    'The Stock page provides real-time inventory levels across all warehouses and depots.'
)

add_bullet('Summary cards show: Total Products in stock, Batch count, and Warehouse locations')
add_bullet('Filter by warehouse/depot or product category')
add_bullet('Low stock items are highlighted with warning indicators')
add_bullet('View batch numbers and expiry dates')
add_bullet('Export stock summary to Excel')

doc.add_heading('Low Stock Alerts', level=3)
doc.add_paragraph(
    'Products that fall below the configured reorder point will appear with a warning badge. '
    'You will also receive a notification in the Notifications page.'
)

# 4.3 Prices
doc.add_heading('4.3 Prices', level=2)
doc.add_paragraph(
    'The Prices page (Admin only) allows you to view and manage price lists synced from SAP.'
)

add_bullet('Select a price list from the dropdown to view all item prices')
add_bullet('Each price list shows: Name, Currency, and Price Factor')
add_bullet('Search for specific items by code or description')
add_bullet('Click "Refresh" to pull the latest prices from SAP')

# 4.4 Inventory Transfers
doc.add_heading('4.4 Inventory Transfers', level=2)
doc.add_paragraph(
    'Inventory Transfers allow you to move stock between warehouses or depots.'
)

doc.add_heading('Creating a Transfer Request', level=3)
add_numbered('Navigate to Inventory → Create Transfer Request.')
add_numbered('Select the source warehouse and destination warehouse.')
add_numbered('Add products and quantities to transfer.')
add_numbered('Add a reason and any notes.')
add_numbered('Review the summary and click "Submit".')

doc.add_heading('Creating an Inventory Transfer', level=3)
add_numbered('Navigate to Inventory → Create Inventory Transfer.')
add_numbered('Follow the step-by-step wizard: From Warehouse → To Warehouse → Items → Review.')
add_numbered('The system validates that sufficient stock is available at the source.')
add_numbered('Submit the transfer to post it to SAP.')

doc.add_heading('Transfer Statuses', level=3)
add_table(
    ['Status', 'Meaning'],
    [
        ['Requested', 'Transfer request submitted, awaiting processing'],
        ['In Transit', 'Goods have left the source warehouse'],
        ['Received', 'Goods received at the destination warehouse'],
        ['Cancelled', 'Transfer was cancelled'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 5. PURCHASING
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('5. Purchasing', level=1)

doc.add_heading('5.1 Purchase Orders', level=2)
doc.add_paragraph(
    'The Purchase Orders page (Admin and Manager only) lets you create and track '
    'orders placed with suppliers.'
)

doc.add_heading('Viewing Purchase Orders', level=3)
add_bullet('Summary cards show: Total POs, Open, On Order, and Received')
add_bullet('Filter by supplier, date range, or status')
add_bullet('Search by PO number or supplier name')

doc.add_heading('Creating a Purchase Order', level=3)
add_numbered('Navigate to Purchasing → Create PO.')
add_numbered('Select the supplier from the vendor list.')
add_numbered('Add items to order with quantities and expected prices.')
add_numbered('Set the expected delivery date.')
add_numbered('Select payment terms and currency.')
add_numbered('Add freight charges and special instructions if needed.')
add_numbered('Click "Save" to post to SAP.')

add_table(
    ['Status', 'Meaning'],
    [
        ['Draft', 'PO being prepared, not yet sent'],
        ['Sent', 'PO sent to the supplier'],
        ['Confirmed', 'Supplier confirmed the order'],
        ['Partially Received', 'Some goods received, others pending'],
        ['Received', 'All goods received'],
        ['Cancelled', 'PO was cancelled'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 6. PAYMENTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('6. Payments', level=1)

doc.add_heading('6.1 Viewing Payments', level=2)
doc.add_paragraph(
    'The Payments page shows all incoming customer payments with summary statistics.'
)
add_bullet('Summary cards show: Total Payments, Today\'s Total, Outstanding, and Latest Payment')
add_bullet('Filter by customer, date range, or payment method (Cash, Cheque, Bank Transfer, Card)')
add_bullet('Search by customer name or payment reference')
add_bullet('Click any payment to view details and the invoices it was applied to')

doc.add_heading('6.2 Recording a Payment', level=2)
doc.add_paragraph('To record a new incoming payment:')
add_numbered('Navigate to Payments → Create Payment.')
add_numbered('Select the customer.')
add_numbered('Choose the payment method (Cash, Cheque, Bank Transfer, Card, or Credit).')
add_numbered('Enter the payment reference or cheque number.')
add_numbered('Enter the payment amount and date.')
add_numbered('Select the invoices to apply the payment against. You can make partial payments.')
add_numbered('Add any notes (e.g., bank name, reference details).')
add_numbered('Click "Save & Post" to record the payment in SAP.')

add_tip('You can apply a single payment across multiple invoices by selecting them from the list.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 7. CUSTOMERS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('7. Customers', level=1)

doc.add_heading('7.1 Customer List', level=2)
doc.add_paragraph(
    'The Customers page is your customer master database, synced from SAP Business One.'
)
add_bullet('Filter by status (Active, Inactive, Blocked), customer type, or region')
add_bullet('Search by customer code, name, or contact details')
add_bullet('Click any customer to view full details: code, name, address, credit limit, and contacts')
add_bullet('View customer transaction history')
add_bullet('Manage credit limits (Admin only)')

doc.add_heading('7.2 G/L Accounts', level=2)
doc.add_paragraph(
    'The G/L Accounts page shows the Chart of Accounts synced from SAP.'
)
add_bullet('Summary cards show: Total Accounts, Revenue, Expenses, and Other accounts')
add_bullet('Filter by account type (Revenue, Expenses, Assets, Liabilities) or currency')
add_bullet('Search by account code or description')
add_bullet('Click "Sync from SAP" to force a refresh of account data')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 8. INSIGHTS & REPORTING
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('8. Insights & Reporting', level=1)

doc.add_heading('8.1 Reports', level=2)
doc.add_paragraph(
    'The Reports page is a comprehensive analytics hub with multiple report tabs:'
)

add_table(
    ['Report', 'Description'],
    [
        ['Sales Summary', 'Revenue trends, quantities sold, and sales performance'],
        ['Top Products', 'Best-selling products by revenue and quantity'],
        ['Stock Summary', 'Current inventory levels, values, and aging'],
        ['Payments Report', 'Collections, outstanding balances, and aging'],
        ['Top Customers', 'Highest-value customers by revenue and frequency'],
        ['Low Stock Alerts', 'Products below reorder point'],
        ['Order Fulfilment', 'Delivery status tracking across all orders'],
        ['Credit Notes', 'Issued, applied, and pending credits'],
        ['Purchase Orders', 'PO status, expenses, and supplier performance'],
        ['Receivables Aging', 'Overdue analysis by age bracket'],
        ['Profit Overview', 'Gross margin and profitability metrics'],
        ['Slow Moving Products', 'Products with no recent activity'],
    ]
)

doc.add_heading('Exporting Reports', level=3)
add_bullet('Export to Excel – download the current report tab or all tabs')
add_bullet('Export to PDF – generate a PDF of the currently selected report')
add_tip('Set a date range before exporting to limit the data to the period you need.')

doc.add_heading('8.2 Proof of Delivery (POD)', level=2)
doc.add_paragraph(
    'The POD page allows operators to upload proof of delivery attachments for invoices.'
)

doc.add_heading('Uploading a POD', level=3)
add_numbered('Navigate to Insights → Proof of Delivery.')
add_numbered('Enter the Invoice Number and click "Lookup" to fetch the invoice details.')
add_numbered('Verify the invoice details (customer, amount, date).')
add_numbered('Click the file upload area or drag-and-drop your POD file (photo, signature, or PDF).')
add_numbered('Optionally add a description.')
add_numbered('Click "Upload" to attach the file.')

add_note('Supported file types: JPG, PNG, WebP, PDF. Maximum file size: 20 MB.')

doc.add_heading('8.3 POD Dashboard', level=2)
doc.add_paragraph(
    'The POD Dashboard shows your personal upload statistics as a POD operator:'
)
add_bullet('Uploads Today – how many PODs you have uploaded today')
add_bullet('Total Uploads – your cumulative upload count')
add_bullet('Completed Deliveries – invoices with POD attached')
add_bullet('Pending – invoices still awaiting POD upload')
add_bullet('Timeline graph of your recent upload activity')

doc.add_heading('8.4 POD Upload Report', level=2)
doc.add_paragraph(
    'The POD Report provides analytics on POD completion rates across all operators.'
)
add_bullet('Quick date selectors: Today, This Week, This Month, or Custom range')
add_bullet('Summary cards: Total Invoices, Uploaded, Pending, and Pending percentage')
add_bullet('Completion summary table')
add_bullet('Export to Excel for further analysis')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 9. DOCUMENTS & TEMPLATES
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('9. Documents & Templates', level=1)

doc.add_heading('9.1 Document Templates', level=2)
doc.add_paragraph(
    'The Document Templates page lets you manage the HTML templates used for generating '
    'invoices, quotations, sales orders, purchase orders, delivery notes, and credit notes.'
)
add_bullet('View all templates as visual cards')
add_bullet('Filter by document type or status (Active/Inactive)')
add_bullet('Set a default template for each document type')
add_bullet('Create, edit, clone, or delete templates')
add_bullet('Configure paper size (A4, Letter, Legal) and orientation')
add_bullet('Use {{Placeholder}} syntax in templates for dynamic data insertion')

doc.add_heading('9.2 Document Manager', level=2)
doc.add_paragraph(
    'The Document Manager is accessed from within specific documents (invoices, orders, etc.) '
    'and provides the following actions:'
)
add_bullet('Generate PDF – select a template and create a PDF document')
add_bullet('Email Document – send to the customer with customisable subject and message')
add_bullet('Upload Attachments – attach additional files to the document')
add_bullet('View History – see previously generated documents')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 10. SYSTEM TOOLS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('10. System Tools', level=1)

doc.add_heading('10.1 Notifications', level=2)
doc.add_paragraph(
    'The Notifications page shows all system alerts and messages.'
)
add_bullet('Filter tabs: All, Unread, Urgent, System')
add_bullet('Search by keyword')
add_bullet('Filter by category: Low Stock, Payment, Invoice, System')
add_bullet('Mark as read/unread or delete')
add_bullet('Click a notification to navigate to the related document')

add_tip('Check the notification bell in the top bar for a quick preview of unread alerts without leaving your current page.')

doc.add_heading('10.2 Exchange Rates', level=2)
doc.add_paragraph(
    'The Exchange Rates page allows you to convert between currencies and manage exchange rates.'
)
add_bullet('Quick Converter: enter an amount, select From and To currencies, and click Convert')
add_bullet('Supported currencies: USD, ZIG, EUR, GBP, ZAR')
add_bullet('View current rate table')
add_bullet('Add new exchange rates or fetch external rates')

doc.add_heading('10.3 Sync Status', level=2)
doc.add_paragraph(
    'The Sync Status page monitors the health of the SAP Business One connection and '
    'data synchronisation.'
)
add_bullet('SAP Connection Status – shows Connected or Disconnected')
add_bullet('System Health – Healthy, Degraded, or Critical')
add_bullet('Pending Queue – number of items waiting to sync')
add_bullet('Health Score – overall system health percentage')
add_bullet('Use "Test SAP Connection" to verify connectivity')
add_bullet('Use "Manual Sync" to trigger an immediate data refresh')

doc.add_heading('10.4 REVMax Fiscal Device', level=2)
doc.add_paragraph(
    'The REVMax page (Admin only) monitors the fiscal device used for tax compliance in Zimbabwe.'
)
add_bullet('Device Card – company TIN, BPN, VAT number, serial number, and device status')
add_bullet('Fiscal Day Card – current fiscal day status (Open/Closed), last receipt number')
add_bullet('License Card – license status and expiry date')
add_bullet('Invoice Lookup – search for a specific invoice\'s fiscal record')
add_bullet('Click "Refresh All" to update all device information')

doc.add_heading('10.5 AI Assistant', level=2)
doc.add_paragraph(
    'The AI Assistant provides a chat interface for querying your business data using '
    'natural language.'
)
add_bullet('Ask questions like "What were our top 5 products last month?"')
add_bullet('Get sales forecasts and trend analysis')
add_bullet('Identify slow-moving stock or unusual transactions')
add_bullet('Generate quick report summaries')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 11. ADMINISTRATION
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('11. Administration (Admin Only)', level=1)
doc.add_paragraph(
    'The following pages are only accessible to users with the Admin role.'
)

doc.add_heading('11.1 User Management', level=2)
doc.add_paragraph(
    'Manage all system user accounts from the User Management page.'
)

doc.add_heading('Users Tab', level=3)
add_bullet('View all users with their role, status, and last login date')
add_bullet('Add User – create a new staff account with role assignment')
add_bullet('Add Mobile User – create a PodOperator account for delivery staff')
add_bullet('Edit user details, reset passwords, lock/unlock, or deactivate accounts')

doc.add_heading('Permissions Tab', level=3)
add_bullet('View the role-based permissions matrix')
add_bullet('Assign or modify roles for each user')

doc.add_heading('Security Tab', level=3)
add_bullet('Configure two-factor authentication requirements')
add_bullet('Set password policy rules')
add_bullet('Configure session timeout duration')

doc.add_heading('11.2 Security Settings', level=2)
doc.add_paragraph(
    'The Security page allows any user to manage their own security settings.'
)
add_bullet('Change your password')
add_bullet('Set up Two-Factor Authentication (Email, SMS, or Authenticator App)')
add_bullet('View active sessions and terminate sessions on other devices')
add_bullet('View your login history and any failed login attempts')

doc.add_heading('11.3 Audit Trail', level=2)
doc.add_paragraph(
    'The Audit Trail provides a complete log of all actions performed in the system.'
)
add_bullet('Summary cards: Total Events, Unique Users, Action Types, Failed Actions')
add_bullet('Filter by date range, user, action type, document type, or status (success/failure)')
add_bullet('Action types include: Create, Edit, Delete, Approve, Reverse, View, Export, Login, Logout')
add_bullet('Export the audit log to CSV')

doc.add_heading('11.4 User Activity', level=2)
doc.add_paragraph(
    'Track user login patterns and system usage on the User Activity page.'
)
add_bullet('Quick time filters: Today, This Week, This Month, or Custom')
add_bullet('Summary cards: Active Users, Total Logins, Average Session Duration, Peak Hour')
add_bullet('View individual user activity breakdowns')
add_bullet('Export to Excel')

doc.add_heading('11.5 Backups', level=2)
doc.add_paragraph('Manage database backups from the Backups page.')
add_bullet('Click "Create Backup" to generate a new database backup')
add_bullet('Summary cards: Total Backups, Successful, Failed, and Total Size')
add_bullet('Download backup files for off-site storage')
add_bullet('Restore from a previous backup (use with caution)')
add_bullet('Configure automated backup schedules and retention policies')

add_note('Restoring a backup will overwrite current data. Always confirm with your team before proceeding.')

doc.add_heading('11.6 Settings', level=2)
doc.add_paragraph('The Settings page provides global application configuration.')

doc.add_heading('Configuration', level=3)
add_bullet('General – Company name, logo, and address')
add_bullet('Display – Date format, number format, timezone, and theme')
add_bullet('Printing – Default printer, paper size, and document templates')

doc.add_heading('Integration', level=3)
add_bullet('Payments – Payment gateway credentials and merchant codes')
add_bullet('SAP – Connection settings and sync frequency')
add_bullet('Email – SMTP server, sender address, and credentials')

doc.add_heading('11.7 Customer Portal Management', level=2)
doc.add_paragraph(
    'Manage which customers have access to the external Customer Portal.'
)
add_bullet('Summary cards: Active Portals, Linked Accounts, Pending Activations')
add_bullet('Enable or disable portal access per customer')
add_bullet('Link multi-account customers')
add_bullet('Configure portal permissions (Invoices, Statements, Payments)')
add_bullet('Reset customer portal passwords')

doc.add_heading('11.8 Webhooks', level=2)
doc.add_paragraph(
    'Configure HTTP webhooks to receive real-time notifications when events '
    'occur in the system (e.g., invoice created, payment received, stock updated).'
)
add_bullet('Create a webhook by specifying the target URL, events, and secret key')
add_bullet('Test webhooks with a sample payload')
add_bullet('View delivery history and HTTP status codes')
add_bullet('Edit or delete existing webhooks')

doc.add_heading('11.9 API Explorer', level=2)
doc.add_paragraph(
    'The API Explorer and Swagger pages provide interactive documentation for the '
    'Kefalos Portal REST API.'
)
add_bullet('Browse API modules by category')
add_bullet('Test API calls directly from the browser')
add_bullet('View request/response examples and authentication requirements')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 12. CUSTOMER PORTAL
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('12. Customer Portal (External)', level=1)
doc.add_paragraph(
    'The Customer Portal is a self-service area for your customers to view their account '
    'information, download invoices, check statements, and track deliveries. Access is '
    'granted by an administrator through the Customer Portal Management page.'
)

doc.add_heading('12.1 Customer Login', level=2)
doc.add_paragraph('Customers access the portal via a separate login page:')
add_numbered('Navigate to the Customer Portal URL.')
add_numbered('Enter the customer code or email address and password.')
add_numbered('If 2FA is enabled, enter the verification code.')
add_numbered('Click "Sign In" to access the dashboard.')

doc.add_heading('12.2 Customer Dashboard', level=2)
add_bullet('Welcome message with customer name and code')
add_bullet('Account Balance indicator (positive or negative)')
add_bullet('Outstanding invoices count and overdue amounts')
add_bullet('Last payment date and amount')
add_bullet('Quick-action buttons: View Statement, View Invoices')
add_bullet('Recent transactions list (last 10)')

doc.add_heading('12.3 Customer Invoices', level=2)
add_bullet('View all invoices with status: Open, Partially Paid, Paid, Overdue')
add_bullet('Filter by date range or status')
add_bullet('Search by invoice number')
add_bullet('Download invoice PDF')
add_bullet('View line item details')

doc.add_heading('12.4 Customer Statements', level=2)
add_bullet('Select a period: Month, Quarter, Year, or Custom Date Range')
add_bullet('View account statement with running balance')
add_bullet('Aging analysis: 0–30, 31–60, 61–90, and 90+ days')
add_bullet('Download statement as PDF or email it')

doc.add_heading('12.5 Customer Payments', level=2)
add_bullet('View complete payment history with dates, amounts, and references')
add_bullet('Download payment receipts')
add_bullet('See current account balance and overdue amounts')

doc.add_heading('12.6 Customer PODs', level=2)
add_bullet('View delivery status for all invoices')
add_bullet('Download proof of delivery attachments (photos, signatures)')
add_bullet('Filter by date range')

doc.add_heading('12.7 Customer Support', level=2)
add_bullet('View support contact information')
add_bullet('Submit support tickets and track their status')
add_bullet('Attach documentation to tickets')
add_bullet('View FAQ section')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 13. TROUBLESHOOTING & FAQ
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('13. Troubleshooting & FAQ', level=1)

doc.add_heading('Common Issues', level=2)

doc.add_heading('I cannot log in', level=3)
add_bullet('Verify your username and password are correct (passwords are case-sensitive)')
add_bullet('Check if your account has been locked after too many failed attempts')
add_bullet('Contact your administrator to reset your password or unlock your account')

doc.add_heading('The page shows "Disconnected" or a loading spinner', level=3)
add_bullet('Check your network/internet connection')
add_bullet('Try refreshing the page (Ctrl+Shift+R for a hard refresh)')
add_bullet('Check the Sync Status page – the SAP connection may be down')
add_bullet('Contact your administrator if the issue persists')

doc.add_heading('I cannot see certain menu items', level=3)
add_bullet('Menu visibility is based on your assigned role')
add_bullet('Contact your administrator if you need access to additional features')

doc.add_heading('An invoice or order is not appearing', level=3)
add_bullet('Allow a few seconds for SAP synchronisation to complete')
add_bullet('Use the "Refresh" button on the page')
add_bullet('Check the Sync Status page for any pending or failed sync items')

doc.add_heading('Export to Excel is not working', level=3)
add_bullet('Ensure your browser allows file downloads')
add_bullet('Check that you have applied filters – exporting very large datasets may take longer')

doc.add_heading('POD upload fails', level=3)
add_bullet('Check the file type – only JPG, PNG, WebP, and PDF are supported')
add_bullet('Check the file size – maximum 20 MB per file')
add_bullet('Ensure the invoice number is valid and exists in SAP')

doc.add_heading('Frequently Asked Questions', level=2)

p = doc.add_paragraph()
r = p.add_run('Q: Can I use the system on my phone?')
r.bold = True
doc.add_paragraph('A: Yes. The Kefalos Portal is fully responsive and works on mobile browsers. The POD Dashboard and upload features are optimised for field use on mobile devices.')

p = doc.add_paragraph()
r = p.add_run('Q: How often does data sync with SAP?')
r.bold = True
doc.add_paragraph('A: Data is synchronised in real time when you create or modify documents. Background syncs for stock and pricing run periodically. You can check the Sync Status page for details.')

p = doc.add_paragraph()
r = p.add_run('Q: Can I undo a posted invoice?')
r.bold = True
doc.add_paragraph('A: Posted invoices cannot be deleted. Instead, create a Credit Note to reverse or adjust the invoice amount.')

p = doc.add_paragraph()
r = p.add_run('Q: What is the REVMax fiscal device?')
r.bold = True
doc.add_paragraph('A: REVMax is the fiscal device required for tax compliance in Zimbabwe. It generates fiscal receipts with QR codes for transactions. The Kefalos Portal automatically fiscalises invoices and credit notes through this device.')

p = doc.add_paragraph()
r = p.add_run('Q: How do I switch to dark mode?')
r.bold = True
doc.add_paragraph('A: Click the theme toggle icon in the top navigation bar to switch between light and dark mode. Your preference is saved automatically.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# APPENDIX A: ROLES & PERMISSIONS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('Appendix A: User Roles & Permissions', level=1)

add_table(
    ['Role', 'Access Areas', 'Key Capabilities'],
    [
        ['Admin', 'All areas', 'Full system administration, user management, settings, backups, webhooks, API access'],
        ['Cashier', 'Sales, Payments, Inventory, Customers, Reports, PODs', 'Create invoices, quotations, credit notes; record payments; manage PODs'],
        ['Manager', 'Purchasing, Reports, User Activity, Inventory, Customers', 'Create purchase orders; view reports, user activity; manage products'],
        ['StockController', 'Inventory, Transfers, Products, Customers, Reports', 'Monitor stock; create transfer requests; manage inventory'],
        ['DepotController', 'Inventory, Transfers, Payments, Customers, Reports', 'Manage depot operations; record payments; process transfers'],
        ['PodOperator', 'POD Dashboard, Proof of Delivery, POD Report', 'Upload proof of delivery attachments; view personal upload stats'],
        ['Customer', 'Customer Portal only', 'View invoices, statements, payments, PODs; submit support tickets'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# APPENDIX B: KEYBOARD SHORTCUTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('Appendix B: Keyboard Shortcuts', level=1)

add_table(
    ['Shortcut', 'Action'],
    [
        ['Ctrl + F', 'Focus the search/filter bar on the current page'],
        ['Ctrl + P', 'Print the current page or document'],
        ['Ctrl + Shift + R', 'Hard refresh the page (clears cache)'],
        ['Esc', 'Close the current dialog, drawer, or modal'],
        ['Tab / Shift+Tab', 'Navigate between form fields'],
        ['Enter', 'Submit the current form or confirm a dialog'],
    ]
)

# ── Footer note ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— End of Document —')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Kefalos Portal End-User Guide  |  Generated {datetime.now().strftime("%d %B %Y")}')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0xbb, 0xbb, 0xbb)

# ── Save ──
output = r'c:\Users\ngoni\source\repos\Ngonie69\ShopInventory\Kefalos_Portal_User_Guide.docx'
doc.save(output)
print(f'Document saved to: {output}')
